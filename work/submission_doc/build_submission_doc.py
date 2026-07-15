from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = Path(r"C:\Users\ASUS\Downloads\智能背包作品信息_已补全草稿_交通避险版.docx")


FONT_BODY = "Microsoft YaHei"
FONT_HEADING = "Microsoft YaHei"
BLUE = RGBColor(31, 78, 121)
MUTED = RGBColor(90, 103, 116)
TODO_RED = RGBColor(156, 49, 42)


def set_run_font(run, size=None, bold=None, color=None, font=FONT_BODY):
    run.font.name = font
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font)
    run._element.rPr.rFonts.set(qn("w:ascii"), font)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), font)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = color


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_widths(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(sum(widths)))

    tbl_grid = tbl.tblGrid
    if tbl_grid is None:
        tbl_grid = OxmlElement("w:tblGrid")
        tbl.insert(0, tbl_grid)
    for child in list(tbl_grid):
        tbl_grid.remove(child)
    for width in widths:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        tbl_grid.append(grid_col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:type"), "dxa")
            tc_w.set(qn("w:w"), str(widths[idx]))
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def style_doc(doc):
    section = doc.sections[0]
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.header_distance = Inches(0.49)
    section.footer_distance = Inches(0.49)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = FONT_BODY
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_BODY)
    normal._element.rPr.rFonts.set(qn("w:ascii"), FONT_BODY)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT_BODY)
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.15

    for style_name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, RGBColor(31, 77, 120), 8, 4),
    ]:
        st = styles[style_name]
        st.font.name = FONT_HEADING
        st._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_HEADING)
        st._element.rPr.rFonts.set(qn("w:ascii"), FONT_HEADING)
        st._element.rPr.rFonts.set(qn("w:hAnsi"), FONT_HEADING)
        st.font.size = Pt(size)
        st.font.color.rgb = color
        st.font.bold = True
        st.paragraph_format.space_before = Pt(before)
        st.paragraph_format.space_after = Pt(after)
        st.paragraph_format.keep_with_next = True

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("SS928-SmartBag 作品信息草稿")
    set_run_font(run, size=8, color=MUTED)


def add_para(doc, text="", style=None, bold=False, color=None, align=None):
    p = doc.add_paragraph(style=style)
    if align is not None:
        p.alignment = align
    r = p.add_run(text)
    set_run_font(r, bold=bold, color=color)
    return p


def add_todo(doc, text):
    p = add_para(doc)
    r1 = p.add_run("待补充/待确认：")
    set_run_font(r1, bold=True, color=TODO_RED)
    r2 = p.add_run(text)
    set_run_font(r2, color=TODO_RED)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        r = p.add_run(item)
        set_run_font(r)


def add_numbers(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        r = p.add_run(item)
        set_run_font(r)


def add_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        shade_cell(hdr[i], "F2F4F7")
        p = hdr[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        set_run_font(r, bold=True)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            p = cells[i].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(value)
            set_run_font(r, size=9.5)
    if widths:
        set_table_widths(table, widths)
    else:
        set_table_widths(table, [9360 // len(headers)] * len(headers))
    add_para(doc, "")
    return table


def add_connection_logic(doc):
    add_para(doc, "服务器优先连接逻辑", "Heading 3")
    add_para(
        doc,
        "正式系统采用“板端上报服务器、小程序从服务器拉取”的主链路，BLE 只作为无网络、现场调试或服务器不可用时的备用链路。这样可以避免手机必须一直在背包附近，也便于做历史曲线、远程告警、多人查看和后续数据分析。",
    )
    add_numbers(
        doc,
        [
            "SS928 边缘端以固定采样率读取 BMI270，完成姿态解算、异常判断和本地告警帧封装。",
            "边缘端通过以太网、Wi-Fi 或 4G/5G 模块把最新帧与告警事件上传到服务器；推荐 HTTPS REST 或 MQTT，两者都保留设备 ID、时间戳和签名字段。",
            "服务器保存 latest、history、alerts 三类数据，维护设备在线状态，并提供小程序访问接口。",
            "小程序默认登录服务器，按设备 ID 拉取最新姿态帧、告警列表和历史曲线；需要实时性时可用短轮询、长轮询或 WebSocket，但用户侧逻辑仍以服务器为数据源。",
            "当服务器不可达或现场调试时，小程序才进入 BLE 备用模式，扫描 BMI270-Backpack，订阅 Nordic UART Service 的 TX notify，并可写入 ZERO、STATUS 等简单命令。",
        ],
    )
    add_table(
        doc,
        ["链路", "用途", "状态写法"],
        [
            ["服务器主链路", "日常使用、远程查看、历史记录、告警留痕", "目标架构；需补服务端接口、域名、数据库和鉴权方式"],
            ["BLE 备用链路", "无网调试、现场近距离连接、服务器故障时应急查看", "当前代码已实现 BLE NUS 原型，可作为备用能力描述"],
            ["本地震动/语音提醒", "主要提醒方式，用于让儿童穿戴者在交通风险出现时及时躲避", "历史记忆显示 MAX98357 播放路径已验证；震动电机和联动代码待补"],
        ],
        [1800, 3300, 4260],
    )


def add_warning_fusion_logic(doc):
    add_para(doc, "雷达 + 摄像头双重预警算法", "Heading 3")
    add_para(
        doc,
        "主程序的预警不应写成单纯由摄像头 YOLO 测速决定。正式逻辑是双通道融合：雷达测速先形成一层风险判断，摄像头/YOLO 视觉测速再形成一层风险判断，融合模块再输出最终 1-4 级预警。这样既利用雷达对速度和距离变化敏感的优势，也利用摄像头对目标类别、画面位置和视觉轨迹的识别能力，减少单一传感器误判。",
    )
    add_table(
        doc,
        ["层级", "输入", "输出", "说明"],
        [
            ["雷达测速层", "目标距离、速度、来向、状态位、轨迹持续时间", "radar_level 1-4 或无风险", "用于判断快速靠近、距离过近、持续高速等动态风险"],
            ["摄像头测速层", "YOLO 目标框、类别、帧间位置变化、视觉估速", "camera_level 1-4 或无风险", "用于判断目标类型、画面区域、视觉速度和运动趋势"],
            ["融合预警层", "radar_level、camera_level、持续时间、一致性", "final_warning_level 1-4", "两路信息结合得到最终预警等级，而不是只采用摄像头结果"],
        ],
        [1800, 3000, 2100, 2460],
    )
    add_para(
        doc,
        "融合原则建议写为：两路各自独立给出子等级；若雷达和摄像头都检测到同一方向的高速/靠近风险，则采用更高等级或上调等级；若只有一路触发，则结合持续时间和可信度保守输出；若两路冲突，则进入短时间复核，避免单帧误报。最终 1-4 级预警中，3 级和 4 级作为后续摔倒报警逻辑的高风险前置条件。",
    )
    add_todo(doc, "补充雷达速度阈值、视觉测速阈值、融合权重或取高/升降级规则，以及 1-4 级每一级的准确含义。")


def add_alarm_escalation_logic(doc):
    add_para(doc, "强避险提醒与家长辅助通知逻辑", "Heading 3")
    add_para(
        doc,
        "报警不是普通姿态异常立即触发，也不是预警后只通知家长。产品的主要目的是减少交通事故，尤其保护儿童穿戴者，让孩子在危险接近时能及时躲避。板端 IMU 摔倒检测状态机先判断是否出现 possible_fall、impact、posture_changed、fall_confirmed 等事件；同时系统回看最近一段时间内是否出现过 3 级或 4 级最终交通风险预警。只有在“摔倒检测成立 + 最近存在 3/4 级风险预警 + 佩戴者 10 秒内没有爬起”同时满足时，板端才发送强提醒/报警信号。该信号的第一作用是驱动背包侧震动和语音提醒穿戴者避险或求助；小程序接收到事件后，再作为辅助链路向家长发送信息提醒。",
    )
    add_table(
        doc,
        ["步骤", "判断内容", "输出"],
        [
            ["1. 风险预警缓存", "雷达 + 摄像头融合输出 1-4 级预警，系统保留最近一段时间的等级记录", "recent_high_risk=true/false"],
            ["2. 摔倒检测", "IMU 跌倒/冲击状态机检测 possible_fall、impact、posture_changed、fall_confirmed", "fall_confirmed 或 impact_only"],
            ["3. 10 秒未爬起判断", "摔倒后观察姿态是否恢复站立/正常运动；若 10 秒内未恢复，认为需要报警", "not_recovered=true/false"],
            ["4. 板端强提醒信号", "满足 fall_confirmed、recent_high_risk、not_recovered 三条件", "驱动震动/语音提醒穿戴者，同时发送 alarm_event 到服务器/小程序备用 BLE"],
            ["5. 家长辅助通知", "小程序或小程序云端接收 alarm_event", "向家长发送信息提醒，并在界面突出显示，但不是唯一提醒方式"],
        ],
        [1300, 5200, 2860],
    )
    add_todo(doc, "确认“最近一段时间”的窗口长度、3/4 级风险的来源字段、10 秒未爬起的姿态判定阈值、震动电机控制方式、语音提示文案，以及微信家长通知采用订阅消息、云函数还是应用内通知。")


def build():
    doc = Document()
    style_doc(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("作品名称")
    set_run_font(r, size=14, bold=True, color=BLUE)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("SS928-SmartBag：基于多传感融合的儿童交通安全避险提醒系统")
    set_run_font(r, size=18, bold=True, color=RGBColor(11, 37, 69))
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("暂定名称，正式提交前可替换为最终作品名；本文未写入学校名称、指导老师等信息。")
    set_run_font(r, size=9, color=MUTED)

    add_para(doc, "资料状态说明", "Heading 2")
    add_para(
        doc,
        "本稿根据当前 SS928 工作区文件、项目记忆和最新口径整理。可以确认的内容直接写入；证据不足、当前仓库缺文件、需要照片或实测数据支撑的内容，均以“待补充/待确认”标注。正式提交前建议统一替换所有待补标记。",
    )

    add_para(doc, "摘要", "Heading 1")
    add_para(
        doc,
        "本作品面向儿童和学生通勤场景下的交通事故预防需求，构建了一套以 SS928/HiEuler Pi 开发板为边缘计算节点、雷达与摄像头为外部交通风险预警输入、BMI270 六轴 IMU 为佩戴者姿态与摔倒感知单元、服务器和微信小程序为数据同步与辅助通知端的智能背包原型系统。产品的主要目的不是单纯记录数据，也不是预警后只通知家长，而是在车辆或高速目标接近时，第一时间通过背包侧震动和语音提醒儿童穿戴者躲避风险，从而减少交通事故。主程序的预警算法不是单一摄像头 YOLO 测速，而是雷达测速层和摄像头/YOLO 视觉测速层双重判断后融合输出 1-4 级最终预警，其中 3 级和 4 级作为高风险前置条件。报警链路在此基础上进一步结合 IMU 摔倒检测：当板端摔倒检测成立、最近一段时间内存在 3/4 级风险预警，并且佩戴者 10 秒内没有爬起时，系统才发送强提醒/报警信号，优先驱动震动与语音提醒穿戴者，随后同步服务器和小程序，由小程序向家长发送辅助信息提醒。理想工作模式下，板端通过网络把数据和告警上传服务器，小程序从服务器拉取最新状态、历史记录和告警信息；BLE 直连仅作为无网络或现场调试时的备用链路。当前工作区已经落地 BMI270 I2C 读取、BLE NUS 姿态仪表盘、雷达串口解析原型和摄像头 MIPI 预览样例；项目记忆中还记录了 IMU 摔倒检测状态机、CSV 姿态阈值调试、MAX98357 I2S 语音提醒路径和语音文件部署验证。作品的核心价值在于把“外部交通风险预警、穿戴者即时避险提醒、摔倒后复核、服务器留痕、家长辅助知晓”组合成分层、可复核的安全闭环。",
    )
    add_todo(doc, "补充最终服务器是否已经部署、服务器地址/API、实物照片、界面截图和最终量化测试结果。")

    add_para(doc, "第一部分  作品概述", "Heading 1")

    add_para(doc, "功能与特性", "Heading 2")
    add_bullets(
        doc,
        [
            "双重预警：雷达测速层和摄像头/YOLO 视觉测速层分别判断风险，再融合得到最终 1-4 级预警。",
            "姿态与摔倒感知：使用 BMI270 采集三轴加速度和三轴角速度，板端计算姿态角、冲击、姿态变化和未爬起状态。",
            "穿戴者避险提醒：3/4 级交通风险优先触发震动和语音，让儿童穿戴者及时躲避；家长通知是同步辅助链路。",
            "报警复核：摔倒检测成立后，还要检查最近是否出现 3/4 级风险预警，并确认 10 秒内未爬起，才发送强提醒/报警信号。",
            "服务器优先：正式连接逻辑为板端上传服务器，小程序从服务器拉取最新数据和告警记录；BLE 仅作为备用连接和调试通道。",
            "移动端显示：微信小程序展示姿态角、运动状态、原始三轴数据、告警文本和最后更新时间。",
            "本地提醒扩展：震动和 MAX98357 I2S 语音播放用于第一时间提醒穿戴者躲避，家长消息用于事后同步和辅助看护。",
        ],
    )

    add_para(doc, "应用领域", "Heading 2")
    add_para(
        doc,
        "该系统首先面向儿童和学生上下学、过路口、骑行或步行通勤等交通事故高风险场景，用于在车辆或高速目标接近时提醒穿戴者主动躲避。它也可扩展到老人外出看护、通勤背包风险预警、户外运动包安全提示、物流或巡检人员负重状态观察，以及智能穿戴和智能行李箱原型验证。雷达和摄像头的双重测速预警适合判断外部目标快速靠近或交通环境风险，BMI270 摔倒检测适合判断佩戴者是否已经发生跌倒或长时间未起身。服务器主链路使数据不局限于手机近距离连接，能够支持家长或看护者远程查看与辅助接收提醒；BLE 备用链路则适合现场安装、调试和离线演示。",
    )

    add_para(doc, "主要技术特点", "Heading 2")
    add_bullets(
        doc,
        [
            "多传感融合：雷达速度层、摄像头/YOLO 视觉速度层、BMI270 摔倒检测层分工明确，最终报警经过多条件复核。",
            "边缘计算：测速预警、姿态/摔倒判断和即时震动/语音提醒在 SS928 板端完成，手机和服务器主要负责展示、存储与辅助通知。",
            "云端主链路：采用设备上报、服务器保存、小程序拉取的连接方式，避免手机必须一直与背包保持蓝牙连接。",
            "可降级通信：服务器不可用时保留 BLE NUS 直连，保障现场调试和无网演示。",
            "可校准阈值：雷达速度阈值、视觉测速阈值、IMU 摔倒阈值和未爬起时间窗都可以通过实测数据回放修正。",
            "模块化扩展：BMI270、服务器接口、小程序、MAX98357 提醒、网络模块均可独立替换或升级。",
        ],
    )

    add_para(doc, "主要性能指标", "Heading 2")
    add_table(
        doc,
        ["指标", "当前/目标值", "依据或状态"],
        [
            ["IMU 采样频率", "50 Hz 默认配置", "work/linux_bmi270_backpack/config.ss928_ble.json"],
            ["BLE 备用上报频率", "10 Hz 默认配置", "当前 BLE NUS 原型已落地"],
            ["服务器数据刷新", "待实测；建议 latest 接口 1 s 内更新", "目标架构，需补服务端实现"],
            ["雷达测速输入", "距离、速度、来向、状态位、轨迹持续时间", "work/ld2417-radar-web 解析原型可证；最终主程序待补"],
            ["摄像头视觉输入", "YOLO 目标、目标框轨迹、视觉估速", "SDK/ModelZoo 有 YOLO 资料；最终测速阈值待补"],
            ["融合预警输出", "最终 1-4 级预警，3/4 级参与报警复核", "规则由用户口径确认；具体阈值待补"],
            ["姿态输出字段", "roll/pitch/yaw、speed、ag、gyr、st、a、w、al", "板端 compact JSON 与小程序解析逻辑"],
            ["摔倒检测状态", "NORMAL、POSSIBLE_FALL、IMPACT、POSTURE_CHANGED、FALL_CONFIRMED、IMPACT_ONLY", "记忆可证；当前树需补回 work/imu_fall_detector"],
            ["报警复核窗口", "摔倒 + 最近 3/4 级风险 + 10 s 未爬起", "最新用户口径；窗口和阈值待确认"],
            ["驼背检测参考阈值", "历史回放：corrected pitch < -15.5 deg，持续 3.0 s，gyro <= 30 dps，accel 0.75-1.25 g", "记忆可证；当前仓库需确认是否已合入"],
            ["本地避险提醒", "震动 + 语音；语音文件可采用 48 kHz stereo AAC，文件名 audio_chn0.aac", "震动控制待补；历史 MAX98357 路径可证；当前联动代码待补"],
        ],
        [2200, 3600, 3560],
    )

    add_para(doc, "主要创新点", "Heading 2")
    add_bullets(
        doc,
        [
            "把外部风险预警写成雷达测速与摄像头 YOLO 视觉测速双重融合，而不是单传感器判断。",
            "把预警目的明确为提醒儿童穿戴者及时躲避交通风险，震动和语音是第一提醒方式。",
            "把报警设计成“高等级交通风险 + 摔倒确认 + 10 秒未爬起”的复核链路，家长通知作为辅助链路，减少普通误触发。",
            "把手机蓝牙直连升级为“服务器主链路 + BLE 备用链路”，更适合家长远程查看和历史追踪。",
            "用真实佩戴 CSV 数据反推姿态阈值，并加入持续时间、角速度和加速度门控，降低误报。",
            "将 SS928 边缘计算、震动/语音即时提醒、微信小程序展示和服务器同步整合成可演示的智能背包闭环。",
        ],
    )

    add_para(doc, "设计流程", "Heading 2")
    add_numbers(
        doc,
        [
            "明确背包姿态监测需求，确定 BMI270 + SS928 的边缘采集方案。",
            "确定外部风险预警由雷达测速和摄像头/YOLO 测速双层判断共同组成。",
            "完成 I2C0 接线、BMI270 数据读取、姿态解算和 CSV 输出。",
            "先用 BLE 打通板端到小程序的最小数据闭环，便于现场调试。",
            "采集站直、驼背、走路、弯腰等 labeled CSV，回放确定初版阈值。",
            "按最新目标调整为服务器优先架构，小程序默认从服务器拉取数据，BLE 作为备用。",
            "扩展震动电机、MAX98357 语音提醒和后续多设备、历史曲线、远程辅助告警能力。",
        ],
    )

    add_para(doc, "第二部分  系统组成及功能说明", "Heading 1")
    add_para(doc, "整体介绍", "Heading 2")
    add_para(
        doc,
        "系统由感知层、边缘计算层、通信层、服务器层、小程序展示层和本地提醒层组成。感知层的雷达与摄像头识别外部交通风险，BMI270 固定在背包背部中线上方，采集人体背包姿态变化；SS928 作为边缘节点读取传感器并完成风险融合、姿态解算和强提醒控制；通信层默认使用网络把数据上传服务器，BLE NUS 保留为备用；服务器负责最新数据、历史数据、告警事件和设备状态管理；微信小程序从服务器拉取数据并展示仪表盘，必要时辅助提醒家长；震动电机和 MAX98357 扬声器模块用于第一时间提醒穿戴者躲避。",
    )
    add_para(
        doc,
        "文字版整体框图：雷达 + 摄像头/YOLO -> SS928 融合预警 -> 震动/语音提醒穿戴者避险；BMI270 IMU -> SS928 摔倒检测与未爬起复核 -> 强提醒/报警事件；SS928 -> 服务器 API/数据库 -> 微信小程序；备用支路为 SS928 BLE NUS -> 微信小程序 BLE 调试页；本地提醒支路为 SS928 -> 震动电机/MAX98357 I2S -> 穿戴者。",
    )
    add_todo(doc, "替换为正式系统框图，可用 draw.io/Visio 绘制；补充服务器、数据库、网络模块和小程序截图。")
    add_connection_logic(doc)
    add_warning_fusion_logic(doc)
    add_alarm_escalation_logic(doc)

    add_para(doc, "硬件系统介绍", "Heading 2")
    add_para(doc, "2.2.1 硬件整体介绍", "Heading 3")
    add_table(
        doc,
        ["模块", "作用", "当前状态"],
        [
            ["SS928/HiEuler Pi 开发板", "边缘计算、传感器读取、网络/BLE 通信、告警控制", "当前工作区围绕该板开发"],
            ["雷达模块", "提供目标距离、速度、来向和高速状态，形成雷达测速风险层", "当前有 LD2417 Web 解析原型；最终板端融合待补"],
            ["摄像头模块", "采集图像，配合 YOLO 完成目标检测、轨迹与视觉测速", "当前有 IMX347 MIPI 预览样例和 YOLO 资料；最终算法待补"],
            ["BMI270 六轴 IMU", "采集加速度、角速度，计算背包姿态", "I2C0 接线和读取程序已形成"],
            ["网络通信模块", "把板端数据上传服务器", "待确认使用以太网、Wi-Fi 还是 MT5710/其他 4G/5G 模块"],
            ["服务器", "保存最新帧、历史数据、告警事件，给小程序提供接口", "目标架构，当前仓库未见服务端代码"],
            ["微信小程序", "用户查看姿态、告警、历史记录；备用 BLE 调试", "当前落地 BLE 仪表盘，需改为服务器优先"],
            ["震动电机", "交通风险或报警时给穿戴者直接触觉提醒", "待补型号、驱动方式和接线"],
            ["MAX98357 + 扬声器", "交通风险或报警时播放本地语音，提醒穿戴者躲避或求助", "历史记忆显示播放已验证，当前联动代码待补"],
        ],
        [2200, 3700, 3460],
    )

    add_para(doc, "2.2.2 机械设计介绍", "Heading 3")
    add_para(
        doc,
        "第一版机械安装建议将 BMI270 固定在背包靠近人体背部的中线上方，尽量靠近背部大面积贴合区域，并固定传感器线束，避免采集到背包晃动而不是人体姿态。SS928 板、电源和网络模块可放在背包内部稳定口袋；震动电机应靠近肩带或背部接触区域，扬声器朝外或朝上布置，保证儿童穿戴者能立即感知提醒。",
    )
    add_todo(doc, "补充实物正面、斜 45 度、内部走线、传感器固定位置、震动电机位置和扬声器安装位置照片；如有外壳或支架，补 CAD/手绘图。")

    add_para(doc, "2.2.3 电路各模块介绍", "Heading 3")
    add_table(
        doc,
        ["电路/接口", "连接关系", "说明"],
        [
            ["BMI270 I2C0", "SDA -> Pin3/I2C0_SDA；SCL -> Pin5/I2C0_SCL；VCC -> 3.3V；GND 共地；CSB -> 3.3V；SDO -> GND", "地址为 0x68，当前配置使用 /dev/i2c-0"],
            ["雷达 UART", "历史雷达 bring-up 显示 UART4 可用于串口数据；LD2417 原型解析距离、速度、来向等字段", "最终模块型号、波特率和接线需确认"],
            ["摄像头 MIPI", "IMX347 sensor0 -> VI -> VPSS -> VO/MIPI TX，可作为视觉输入链路参考", "YOLO 测速主程序需补截图/代码"],
            ["震动提醒", "SS928 GPIO/PWM 或外接驱动 -> 震动电机", "用于直接提醒穿戴者，接线和驱动待补"],
            ["MAX98357 I2S", "历史验证：Pin12 -> BCLK；Pin38 -> WS/LRC；Pin40 -> SD_TX/DIN；MCLK 可不接；GND 共地", "用于语音避险提醒，需补最终实物接线图"],
            ["网络链路", "SS928 -> 服务器，承载方式待定", "需确认 Wi-Fi/以太网/4G/5G 模块和供电方案"],
            ["本地提醒输出", "软件可输出 ALERT/ALARM 行、写文件/GPIO 或执行 alert_command", "可与震动电机、语音播放脚本或 GPIO 控制联动"],
        ],
        [2100, 4700, 2560],
    )

    add_para(doc, "软件系统介绍", "Heading 2")
    add_para(doc, "2.3.1 软件整体介绍", "Heading 3")
    add_para(
        doc,
        "软件采用边缘端、服务器端、小程序端三层结构。边缘端主程序由雷达测速层、摄像头/YOLO 视觉测速层、融合预警层、BMI270 摔倒检测层、强提醒复核层和数据上传层组成。雷达和摄像头先融合输出最终 1-4 级交通风险预警；其中高等级风险优先触发本地震动和语音，让穿戴者及时躲避。BMI270 摔倒检测负责判断是否发生摔倒、冲击和姿态变化；强提醒复核层检查近期是否出现 3/4 级预警以及 10 秒内是否未爬起，再决定是否发送 alarm_event。服务器端负责设备鉴权、数据接收、最新状态缓存、历史存储和告警查询；小程序端默认访问服务器接口展示数据，并在需要时辅助提醒家长，只有在服务器或网络不可用时进入 BLE 备用模式。当前工作区已有 BMI270 边缘程序、BLE 小程序原型、雷达解析原型和摄像头预览样例，服务器端代码与最终融合主程序需要补齐或提供截图。",
    )
    add_para(doc, "2.3.2 软件各模块介绍", "Heading 3")
    add_table(
        doc,
        ["模块", "关键输入", "关键输出/功能"],
        [
            ["BMI270 采集模块", "I2C/IIO 原始加速度、角速度", "ImuSample、chip id probe、50 Hz 采样"],
            ["姿态解算模块", "ax/ay/az、gx/gy/gz、dt", "roll_deg、pitch_deg、yaw_deg、speed_mps、accel_g、gyro_dps"],
            ["雷达测速风险层", "雷达目标距离、速度、来向、状态位", "radar_level 1-4，标记快速靠近/距离过近风险"],
            ["摄像头视觉风险层", "YOLO 目标、目标框轨迹、视觉测速", "camera_level 1-4，标记视觉目标速度和区域风险"],
            ["融合预警模块", "radar_level、camera_level、持续时间、一致性", "final_warning_level 1-4"],
            ["摔倒检测模块", "IMU 加速度、角速度、jerk、姿态变化、静止状态", "NORMAL/POSSIBLE_FALL/IMPACT/POSTURE_CHANGED/FALL_CONFIRMED/IMPACT_ONLY"],
            ["强提醒复核模块", "fall_confirmed、recent_high_risk、10s 未爬起", "alarm_event，驱动震动/语音，并发送给服务器/小程序"],
            ["服务器上传模块", "compact frame、alert event、device id", "POST/MQTT 上报 latest/history/alerts；待实现"],
            ["服务器 API 模块", "设备 ID、时间范围、用户身份", "latest、history、alerts、commands 接口；待实现"],
            ["小程序展示模块", "服务器返回的最新帧/历史帧", "姿态仪表盘、告警列表、历史曲线；当前需从 BLE 页面改造"],
            ["BLE 备用模块", "NUS TX notify、RX write", "备用近场查看，支持 ZERO/STATUS 调试命令"],
            ["震动/语音提醒模块", "交通风险等级、告警代码、音频文件 ID", "驱动震动并调用 MAX98357 播放 audio_chn0.aac；待补联动代码"],
        ],
        [2100, 3300, 3960],
    )
    add_para(doc, "建议服务器接口草案", "Heading 3")
    add_table(
        doc,
        ["接口", "方向", "用途"],
        [
            ["POST /api/v1/devices/{id}/frames", "板端 -> 服务器", "上传实时姿态帧和在线心跳"],
            ["POST /api/v1/devices/{id}/alerts", "板端 -> 服务器", "上传异常姿态、撞击、失重等告警事件"],
            ["POST /api/v1/devices/{id}/warnings", "板端 -> 服务器", "上传雷达+摄像头融合后的 1-4 级预警"],
            ["POST /api/v1/devices/{id}/alarm", "板端 -> 服务器", "上传满足摔倒+近期 3/4 级风险+10s 未爬起后的强提醒/报警事件，供家长辅助知晓"],
            ["GET /api/v1/devices/{id}/latest", "小程序 -> 服务器", "拉取最新一帧姿态数据"],
            ["GET /api/v1/devices/{id}/history", "小程序 -> 服务器", "按时间范围拉取历史曲线"],
            ["GET /api/v1/devices/{id}/alerts", "小程序 -> 服务器", "拉取告警记录"],
            ["POST /api/v1/devices/{id}/commands", "小程序 -> 服务器 -> 板端", "下发 ZERO、STATUS、阈值调整等命令"],
        ],
        [3300, 2100, 3960],
    )

    add_para(doc, "第三部分  完成情况及性能参数", "Heading 1")
    add_para(doc, "整体介绍", "Heading 2")
    add_para(
        doc,
        "当前项目已经形成了 BMI270 采集与姿态计算、小程序 BLE 仪表盘、SS928 BLE 启动脚本、雷达串口解析原型、摄像头 MIPI 预览样例和一批板端运行配置；历史记忆中还保留了 IMU 摔倒检测、CSV 阈值回放和 MAX98357 语音部署验证。最新需求将产品定位明确为儿童交通事故预防，通信方式调整为服务器主链路，并明确主程序预警为雷达+摄像头双重融合。因此最终版本应在现有原型基础上补齐服务器上传、服务器存储、小程序拉取、融合预警主程序、震动/语音即时提醒和家长辅助通知展示。",
    )
    add_todo(doc, "补充整机正面、斜 45 度照片；补最终服务器页面/API 调试截图；补小程序服务器拉取版本截图。")

    add_para(doc, "工程成果", "Heading 2")
    add_para(doc, "3.2.1 机械成果", "Heading 3")
    add_para(
        doc,
        "已明确传感器安装位置和背包内板卡布置原则：BMI270 固定在背部中线附近，SS928 与供电模块固定在包内，扬声器外放位置避免被布料完全遮挡。当前缺少最终装配照片和结构件图。",
    )
    add_todo(doc, "上传最终实物照片、内部布线照片、传感器固定照片和必要的结构说明。")

    add_para(doc, "3.2.2 电路成果", "Heading 3")
    add_para(
        doc,
        "BMI270 的 SS928 I2C0 接线方案已经在当前文档和代码配置中固定；MAX98357 40Pin I2S 播放路径在历史记忆中已有可听声验证，可作为语音避险提醒基础；震动电机控制和网络上传链路仍需根据最终硬件补充电路连接与供电说明。",
    )
    add_todo(doc, "补最终电路图、40Pin 占用表、网络模块接线、供电方式和电流/续航测试数据。")

    add_para(doc, "3.2.3 软件成果", "Heading 3")
    add_table(
        doc,
        ["成果", "文件/证据", "状态"],
        [
            ["BMI270 背包姿态程序", "work/linux_bmi270_backpack/bmi270_backpack.py", "当前存在"],
            ["雷达测速解析原型", "work/ld2417-radar-web/parser.js、README.md", "当前存在，用于说明雷达速度层"],
            ["摄像头 MIPI 预览", "work/imx347_mipi_preview", "当前存在，用于说明视觉输入链路"],
            ["YOLO 资料/样例", "SDK SVP/NPU YOLO、ModelZoo YOLO 资料", "资料存在；最终视觉测速主程序待补"],
            ["IMU 摔倒检测模块", "work/imu_fall_detector", "记忆中已实现；当前树缺目录，需补回"],
            ["SS928 BLE 配置", "work/linux_bmi270_backpack/config.ss928_ble.json", "当前存在"],
            ["SS928 启动脚本", "work/linux_bmi270_backpack/start_ss928_ble.sh", "当前存在"],
            ["小程序 BLE 姿态仪表盘", "work/ssminiprogram/miniprogram/pages/index/*", "当前存在，需改服务器优先"],
            ["CSV 阈值分析", "calibration_analysis.md", "记忆中存在；当前树缺文件，需补回"],
            ["MAX98357 语音播放/音频部署", "历史工作流与远端路径记录", "记忆可证；当前树缺本地技能和部署目录"],
            ["服务器上传与小程序拉取", "待建服务端与小程序接口层", "待实现/待补充"],
        ],
        [2700, 3900, 2760],
    )

    add_para(doc, "特性成果", "Heading 2")
    add_table(
        doc,
        ["功能/性能", "已达到或目标描述", "证据状态"],
        [
            ["姿态实时显示", "小程序展示 roll/pitch/yaw、速度、加速度、角速度和告警", "当前 BLE 页面可证"],
            ["双通道预警", "雷达测速层 + 摄像头/YOLO 视觉测速层融合输出 1-4 级", "用户口径已确认；具体阈值待补"],
            ["穿戴者提醒", "高等级风险优先触发震动和语音，提醒儿童及时躲避", "用户口径已确认；震动/语音联动待补"],
            ["家长辅助通知", "摔倒检测成立 + 近期 3/4 级风险 + 10 秒未爬起，板端发送事件，小程序辅助通知家长", "用户口径已确认；消息实现待补"],
            ["BLE 备用连接", "扫描 BMI270-Backpack，订阅 NUS TX notify，RX 写 ZERO/STATUS", "当前代码可证"],
            ["服务器主连接", "小程序从服务器拉取 latest/history/alerts", "最新口径，需实现与测试"],
            ["数据采样", "50 Hz IMU 采样，10 Hz BLE 备用发送", "配置可证"],
            ["驼背阈值", "历史回放选定 -15.5 deg + 3 s + 运动门控", "记忆可证，当前代码需确认"],
            ["语音提醒", "MAX98357 通过 sample_audio 2 播放 audio_chn0.aac", "历史可证，需补联动与现场照片"],
            ["震动提醒", "高等级交通风险和强报警事件触发震动", "待补电机、驱动和实测效果"],
        ],
        [2300, 4200, 2860],
    )

    add_para(doc, "第四部分  总结", "Heading 1")
    add_para(doc, "可扩展之处", "Heading 2")
    add_bullets(
        doc,
        [
            "补齐服务器：实现设备鉴权、数据接收、latest 缓存、历史数据库、告警查询和小程序接口。",
            "完善融合预警：补雷达速度等级、视觉速度等级、融合取高/升降级规则和 1-4 级定义。",
            "完善即时提醒闭环：补震动电机控制、语音提示文案、近期 3/4 级风险窗口、10 秒未爬起判断、alarm_event 格式和家长辅助通知实现。",
            "增强小程序：默认走服务器数据源，增加历史曲线、告警时间线、设备绑定和家庭成员共享。",
            "完善板端联网：根据最终硬件选择 Wi-Fi、以太网或 4G/5G 模块，加入断线重传和本地缓存。",
            "提升算法：引入更多佩戴人群和动作样本，逐步从阈值规则升级到轻量分类模型。",
            "增强硬件：优化背包固定结构、扬声器位置、电池续航、防水抗震和低功耗策略。",
        ],
    )

    add_para(doc, "心得体会", "Heading 2")
    add_para(
        doc,
        "本项目最大的体会是，智能硬件作品不能只把单个传感器读数显示出来，而要回到真实场景：减少交通事故，尤其保护儿童。外部预警如果只写成摄像头 YOLO 测速，会忽略雷达对距离和速度变化的优势；实际方案中雷达先判断一层风险，摄像头再判断一层风险，两者融合得到最终 1-4 级交通风险预警，才能更稳地识别快速靠近或危险场景。预警的第一目标不是通知家长，而是让穿戴者本人立即感知危险并躲避，所以震动和语音是产品最关键的提醒方式。报警也不能因为一次姿态变化就直接发送给家长，而应在摔倒检测成立的基础上，结合近期是否存在 3/4 级高风险预警，并观察 10 秒内是否爬起，满足条件后再由板端发出强提醒/报警信号，同时同步小程序，让家长作为辅助链路及时知晓。最初用 BLE 打通板端到小程序，是为了快速验证 BMI270 数据、姿态角和移动端界面；但真实使用时手机不能始终贴近背包，数据也需要留痕和远程查看，因此服务器主链路更符合最终产品逻辑。后续继续完善服务器、融合阈值、震动/语音联动、实物结构、续航和批量测试后，作品会从“可演示原型”更接近“可实际使用的儿童交通安全智能背包”。",
    )

    add_para(doc, "第五部分  参考文献", "Heading 1")
    refs = [
        "Bosch Sensortec. BMI270 Data Sheet and Application Notes.",
        "BlueZ Project. D-Bus GATT API Documentation.",
        "Nordic Semiconductor. Nordic UART Service profile reference.",
        "微信开放文档. 微信小程序蓝牙与网络请求相关 API 文档.",
        "HiEuler Pi / SS928V100 开发板产品规格书与 40Pin 引脚复用资料.",
        "HLK-LD2417 串口协议说明及本项目 LD2417 雷达解析原型.",
        "YOLO/ModelZoo 目标检测模型部署资料及本项目摄像头预览样例.",
        "Analog Devices / Maxim Integrated. MAX98357A I2S Class-D Amplifier Data Sheet.",
        "本项目工作区：work/linux_bmi270_backpack、work/ssminiprogram 及相关 README/配置文件.",
    ]
    for idx, ref in enumerate(refs, 1):
        p = doc.add_paragraph()
        r = p.add_run(f"[{idx}] {ref}")
        set_run_font(r)
    add_todo(doc, "正式提交前补全每篇参考文献的版本号、URL/出版信息和访问日期。")

    add_para(doc, "待补充/待确认清单", "Heading 1")
    add_table(
        doc,
        ["优先级", "项目", "说明"],
        [
            ["高", "服务器实现与接口", "补域名/IP、API、数据库、鉴权、板端上传和小程序拉取截图"],
            ["高", "雷达+摄像头融合预警规则", "补 radar_level、camera_level、final_warning_level 的 1-4 级阈值和融合规则"],
            ["高", "震动/语音即时提醒", "补震动电机型号、驱动接线、语音提示文案、触发时序和现场验证"],
            ["高", "摔倒报警复核规则", "补最近 3/4 级风险窗口、10 秒未爬起判断、alarm_event 格式和家长辅助通知方式"],
            ["高", "最终连接逻辑验证", "证明日常使用走服务器，BLE 只在无网/调试时启用"],
            ["高", "实物与界面图片", "正面、斜 45 度、内部走线、传感器位置、小程序服务器页"],
            ["中", "量化测试", "延迟、掉线重连、服务器刷新、续航、误报/漏报统计"],
            ["中", "阈值文件合入", "把历史 HUNCH 阈值和 calibration_analysis.md 补回当前仓库或替换为最新结果"],
            ["中", "语音提醒联动", "补 warning/alarm -> 音频 ID -> sample_audio/MAX98357 播放的控制代码和现场验证"],
            ["低", "机械设计图", "如有外壳、支架、固定夹或背包改造结构，补 CAD/手绘图"],
        ],
        [1200, 2600, 5560],
    )

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
